"""
Multi-format Book Exporter (EPUB, Bilingual EPUB, DOCX, HTML, Markdown, TXT).
Preserves book styling, generates table of contents, and produces publication-ready files.
"""
import os
import re
import zipfile
import tempfile
import shutil
import unicodedata
from typing import Optional, List
from bs4 import BeautifulSoup
import ebooklib
from ebooklib import epub
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.parser import BookProject, BookChapter, BookParagraph


def normalize_text(text: str) -> str:
    """Normalize Unicode to NFC (precomposed) and fix escaped dollar signs."""
    if not text:
        return ""
    text = unicodedata.normalize('NFC', text)
    text = text.replace(r'\$', '$')
    return text


def unicode_math_fallback(latex: str) -> str:
    """Fallback converter that maps LaTeX symbols to clean HTML / Unicode."""
    res = latex
    replacements = [
        (r'\times', '×'), (r'\cdot', '·'), (r'\in', '∈'), (r'\notin', '∉'),
        (r'\subset', '⊂'), (r'\subseteq', '⊆'), (r'\cup', '∪'), (r'\cap', '∩'),
        (r'\forall', '∀'), (r'\exists', '∃'), (r'\rightarrow', '→'), (r'\to', '→'),
        (r'\leftarrow', '←'), (r'\Rightarrow', '⇒'), (r'\Leftarrow', '⇐'),
        (r'\leq', '≤'), (r'\le', '≤'), (r'\geq', '≥'), (r'\ge', '≥'),
        (r'\neq', '≠'), (r'\ne', '≠'), (r'\approx', '≈'), (r'\sim', '∼'),
        (r'\infty', '∞'), (r'\pm', '±'), (r'\alpha', 'α'), (r'\beta', 'β'),
        (r'\gamma', 'γ'), (r'\delta', 'δ'), (r'\epsilon', 'ε'), (r'\theta', 'θ'),
        (r'\lambda', 'λ'), (r'\mu', 'μ'), (r'\sigma', 'σ'), (r'\tau', 'τ'),
        (r'\phi', 'φ'), (r'\omega', 'ω'), (r'\Delta', 'Δ'), (r'\Sigma', 'Σ'),
        (r'\Omega', 'Ω'), (r'\mathbb{R}', 'ℝ'), (r'\mathbb{N}', 'ℕ'), (r'\mathbb{Z}', 'ℤ'),
        (r'\sum', '∑'), (r'\prod', '∏'), (r'\int', '∫'),
    ]
    for k, v in replacements:
        res = res.replace(k, v)
    res = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1 / \2)', res)
    res = re.sub(r'\\sqrt\{([^}]+)\}', r'√(\1)', res)
    res = re.sub(r'\\text\{([^}]+)\}', r'\1', res)
    res = re.sub(r'_\{([^}]+)\}', r'<sub>\1</sub>', res)
    res = re.sub(r'_([a-zA-Z0-9])', r'<sub>\1</sub>', res)
    res = re.sub(r'\^\{([^}]+)\}', r'<sup>\1</sup>', res)
    res = re.sub(r'\^([a-zA-Z0-9])', r'<sup>\1</sup>', res)
    return f'<span class="math-fallback">{res}</span>'


def latex_to_mathml(latex_str: str, display: bool = False) -> str:
    """Converts a LaTeX mathematical string to valid MathML XML."""
    clean_latex = latex_str.strip()
    if not clean_latex:
        return ""
    try:
        import latex2mathml.converter
        mathml = latex2mathml.converter.convert(clean_latex)
        if display and 'display="inline"' in mathml:
            mathml = mathml.replace('display="inline"', 'display="block"')
        return mathml
    except Exception:
        return unicode_math_fallback(clean_latex)


def format_math_in_html(text: str) -> tuple[str, bool]:
    """
    Parses LaTeX formulas from text and converts them to MathML with academic styling.
    Returns (formatted_html, has_math).
    """
    text = normalize_text(text)
    has_math = False

    # 1. Check for standalone equation with number at start: e.g. Attention(...) = ... (1) Rest of paragraph
    eq_match = re.match(r'^(.*?\\frac\{[^}]+\}\{[^}]+\}[^\n]*?)\s*\((\d+)\)\s+([A-Z\u00C0-\u1EF9].*)$', text, re.DOTALL)
    extra_block = ""
    if eq_match:
        eq_part = eq_match.group(1).strip()
        eq_num = eq_match.group(2)
        text = eq_match.group(3).strip()
        has_math = True
        eq_mathml = latex_to_mathml(eq_part, display=True)
        extra_block = f'<div class="math-block math-equation"><div class="math-formula">{eq_mathml}</div><div class="eq-num">({eq_num})</div></div>\n'

    # 2. Display math $$...$$
    def rep_display(m):
        nonlocal has_math
        has_math = True
        return f'<div class="math-block">{latex_to_mathml(m.group(1), display=True)}</div>'
    text = re.sub(r'\$\$([^\$]+)\$\$', rep_display, text)

    # 3. Inline math $...$
    def rep_inline(m):
        nonlocal has_math
        content = m.group(1).strip()
        if re.match(r'^\d+(\.\d+)?$', content):
            return f"${content}$"
        has_math = True
        return latex_to_mathml(content, display=False)
    text = re.sub(r'\$([^\$]+)\$', rep_inline, text)

    # 4. Residual bare LaTeX expressions like \frac{...}{...}
    def rep_bare_frac(m):
        nonlocal has_math
        has_math = True
        return latex_to_mathml(m.group(0), display=False)
    text = re.sub(r'\\frac\{[^{}]*\}\{[^{}]*\}', rep_bare_frac, text)

    final_html = extra_block + text if extra_block else text
    return final_html, has_math


def format_math_for_docx(text: str) -> str:
    """Formats LaTeX math into clean Unicode text suitable for Word documents."""
    text = normalize_text(text)
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1 / \2)', text)
    text = re.sub(r'\\sqrt\{([^}]+)\}', r'√(\1)', text)
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)
    replacements = [
        (r'\times', '×'), (r'\cdot', '·'), (r'\in', '∈'), (r'\notin', '∉'),
        (r'\leq', '≤'), (r'\geq', '≥'), (r'\neq', '≠'), (r'\approx', '≈'),
        (r'\infty', '∞'), (r'\pm', '±'), (r'\alpha', 'α'), (r'\beta', 'β'),
        (r'\gamma', 'γ'), (r'\delta', 'δ'), (r'\epsilon', 'ε'), (r'\theta', 'θ'),
        (r'\lambda', 'λ'), (r'\mu', 'μ'), (r'\sigma', 'σ'), (r'\tau', 'τ'),
        (r'\phi', 'φ'), (r'\omega', 'ω'), (r'\Delta', 'Δ'), (r'\Sigma', 'Σ'),
        (r'\Omega', 'Ω'), (r'\mathbb{R}', 'ℝ'), (r'\mathbb{N}', 'ℕ'), (r'\mathbb{Z}', 'ℤ'),
        (r'\sum', '∑'), (r'\prod', '∏'), (r'\int', '∫'),
    ]
    for k, v in replacements:
        text = text.replace(k, v)
    text = re.sub(r'\$([^\$]+)\$', r'\1', text)
    return text


class BookExporter:
    """Exports translated or bilingual books to various formats."""

    @classmethod
    def export_epub(cls, project: BookProject, output_path: str, bilingual: bool = False) -> str:
        """
        Exports as EPUB.
        If original source was EPUB, updates original XHTML documents in-place to preserve all images and styles!
        Otherwise, builds a clean new EPUB using ebooklib.
        """
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

        if project.source_format == "epub" and os.path.exists(project.source_file_path):
            return cls._export_epub_inplace(project, output_path, bilingual)
        else:
            return cls._export_epub_fresh(project, output_path, bilingual)

    @classmethod
    def _export_epub_inplace(cls, project: BookProject, output_path: str, bilingual: bool = False) -> str:
        """
        Modifies the original EPUB zip archive by substituting translated text in XHTML files.
        Guarantees 100% preservation of images, CSS, fonts, and cover.
        """
        # Map paragraph IDs to their translations
        para_map = {}
        for chap in project.chapters:
            for p in chap.paragraphs:
                para_map[p.id] = p

        # Map document names to chapters
        doc_chap_map = {c.doc_name: c for c in project.chapters if c.doc_name}

        temp_dir = tempfile.mkdtemp(prefix="epub_export_")
        try:
            with zipfile.ZipFile(project.source_file_path, 'r') as zin:
                zin.extractall(temp_dir)

            # Update XHTML files
            for root, dirs, files in os.walk(temp_dir):
                for f in files:
                    ext = os.path.splitext(f)[1].lower()
                    if ext in ('.xhtml', '.html', '.htm'):
                        rel_path = os.path.relpath(os.path.join(root, f), temp_dir).replace('\\', '/')
                        # Check matching document
                        target_chap = None
                        for chap in project.chapters:
                            if chap.doc_name and (chap.doc_name in rel_path or rel_path.endswith(chap.doc_name)):
                                target_chap = chap
                                break

                        if target_chap:
                            file_full = os.path.join(root, f)
                            with open(file_full, 'r', encoding='utf-8', errors='ignore') as xf:
                                content = xf.read()

                            soup = BeautifulSoup(content, 'html.parser')
                            raw_tags = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'blockquote', 'li', 'div'])
                            tags = [t for t in raw_tags if not (t.name == 'div' and t.find(['p', 'div', 'blockquote']))]

                            for p in target_chap.paragraphs:
                                trans = p.translated_text.strip()
                                if not trans:
                                    continue

                                # Match tag by text
                                for t in tags:
                                    if t.get_text(strip=True) == p.original_text.strip():
                                        if bilingual:
                                            # Bilingual: original in smaller italic or gray, followed by translated
                                            t.clear()
                                            en_span = soup.new_tag("div")
                                            en_span['style'] = "color: #718096; font-size: 0.9em; margin-bottom: 4px; font-style: italic;"
                                            en_span.string = p.original_text
                                            vi_span = soup.new_tag("div")
                                            vi_span.string = trans
                                            t.append(en_span)
                                            t.append(vi_span)
                                        else:
                                            t.string = trans
                                        break

                            with open(file_full, 'w', encoding='utf-8') as xf:
                                xf.write(str(soup))

            # Repack zip as EPUB
            if os.path.exists(output_path):
                os.remove(output_path)

            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                # mimetype must be first and uncompressed per EPUB spec
                mimetype_path = os.path.join(temp_dir, 'mimetype')
                if os.path.exists(mimetype_path):
                    zout.write(mimetype_path, 'mimetype', compress_type=zipfile.ZIP_STORED)

                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        full = os.path.join(root, file)
                        rel = os.path.relpath(full, temp_dir)
                        if rel == 'mimetype':
                            continue
                        zout.write(full, rel)

            return output_path
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    @classmethod
    def _export_epub_fresh(cls, project: BookProject, output_path: str, bilingual: bool = False) -> str:
        """Constructs a new EPUB book from scratch using ebooklib with MathML and Vietnamese font support."""
        book = epub.EpubBook()
        book.set_identifier(f"ai-book-{project.id}")
        book_title = normalize_text(f"{project.title} (Bản Dịch Tiếng Việt)" if not bilingual else f"{project.title} (Song Ngữ Anh - Việt)")
        book.set_title(book_title)
        book.set_language('vi')
        book.add_author(normalize_text(project.author))

        epub_chapters = []
        toc = []

        # Vietnamese-optimized typography with MathML support
        style = '''
        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Times New Roman", "Palatino Linotype", Arial, sans-serif;
            line-height: 1.7;
            margin: 5%;
            color: #1a202c;
            text-rendering: optimizeLegibility;
            -webkit-font-smoothing: antialiased;
        }
        h1, h2, h3 {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            color: #2b6cb0;
            text-align: center;
            font-weight: 700;
        }
        p {
            text-indent: 1.5em;
            margin-bottom: 0.8em;
            text-align: justify;
        }
        .bilingual-en {
            color: #718096;
            font-size: 0.88em;
            font-style: italic;
            margin-bottom: 3px;
            text-indent: 0;
        }
        .bilingual-vi {
            color: #1a202c;
            margin-bottom: 14px;
        }
        .math-block {
            text-align: center;
            margin: 1.4em 0;
            padding: 10px 14px;
            background: #f8fafc;
            border-radius: 6px;
            overflow-x: auto;
        }
        .math-equation {
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        .math-formula {
            display: inline-block;
            margin: 0 auto;
        }
        .eq-num {
            float: right;
            color: #64748b;
            font-size: 0.9em;
        }
        math {
            font-family: "Cambria Math", "Latin Modern Math", "STIX Two Math", serif;
            font-size: 1.05em;
        }
        .math-fallback {
            font-family: "Cambria Math", "Latin Modern Math", "Times New Roman", serif;
            font-style: italic;
        }
        '''
        default_css = epub.EpubItem(uid="style_default", file_name="style/default.css", media_type="text/css", content=style)
        book.add_item(default_css)

        for i, chap in enumerate(project.chapters):
            c_item = epub.EpubHtml(title=normalize_text(chap.title), file_name=f"chap_{i}.xhtml", lang="vi")
            c_item.add_item(default_css)

            html_parts = [f"<h2>{normalize_text(chap.title)}</h2>"]
            chap_has_math = False

            for p in chap.paragraphs:
                if (p.tag == "img" or getattr(p, "image_path", "")) and p.image_path and os.path.exists(p.image_path):
                    img_filename = f"img_{os.path.basename(p.image_path)}"
                    try:
                        with open(p.image_path, "rb") as f_img:
                            epub_img = epub.EpubItem(
                                uid=f"img_{i}_{p.id}",
                                file_name=f"images/{img_filename}",
                                media_type="image/png" if p.image_path.lower().endswith(".png") else "image/jpeg",
                                content=f_img.read()
                            )
                            book.add_item(epub_img)
                            html_parts.append(f'<div style="text-align: center; margin: 18px 0;"><img src="images/{img_filename}" style="max-width: 100%; height: auto;" /></div>')
                    except Exception:
                        pass
                    continue

                trans = p.translated_text.strip() if p.translated_text.strip() else p.original_text
                trans_html, has_m = format_math_in_html(trans)
                if has_m:
                    chap_has_math = True

                if bilingual:
                    orig_html, _ = format_math_in_html(p.original_text)
                    html_parts.append(f'<div class="bilingual-en">{orig_html}</div>')
                    html_parts.append(f'<p class="bilingual-vi">{trans_html}</p>')
                else:
                    tag = p.tag if p.tag in ('h1', 'h2', 'h3', 'blockquote') else 'p'
                    if trans_html.startswith('<div class="math-block'):
                        parts = trans_html.split('\n', 1)
                        if len(parts) == 2 and parts[1].strip():
                            html_parts.append(parts[0])
                            html_parts.append(f'<{tag}>{parts[1].strip()}</{tag}>')
                        else:
                            html_parts.append(trans_html)
                    else:
                        html_parts.append(f'<{tag}>{trans_html}</{tag}>')

            if chap_has_math:
                c_item.properties.append("mathml")

            c_item.content = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="vi" lang="vi">
<head>
  <meta charset="utf-8" />
  <title>{normalize_text(chap.title)}</title>
  <link rel="stylesheet" type="text/css" href="style/default.css" />
</head>
<body>
{''.join(html_parts)}
</body>
</html>""".encode('utf-8')
            book.add_item(c_item)
            epub_chapters.append(c_item)
            toc.append(c_item)

        book.toc = tuple(toc)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ['nav'] + epub_chapters

        epub.write_epub(output_path, book)
        return output_path

    @classmethod
    def export_docx(cls, project: BookProject, output_path: str, bilingual: bool = False) -> str:
        """Exports book as a polished Microsoft Word (.docx) document with clean math typography."""
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        doc = docx.Document()

        # Title Page
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(normalize_text(project.title))
        run.font.name = "Times New Roman"
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59)

        subtitle_p = doc.add_paragraph()
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_text = "Bản dịch Tiếng Việt (AI Literary Edition)" if not bilingual else "Bản dịch Song Ngữ Anh - Việt"
        sub_run = subtitle_p.add_run(normalize_text(sub_text))
        sub_run.font.name = "Times New Roman"
        sub_run.font.size = Pt(14)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(100, 116, 139)

        author_p = doc.add_paragraph()
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_p.add_run(f"Tác giả: {normalize_text(project.author)}")
        author_run.font.name = "Times New Roman"
        author_run.font.size = Pt(12)
        doc.add_page_break()

        for chap in project.chapters:
            # Chapter heading
            heading = doc.add_heading(normalize_text(chap.title), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for p in chap.paragraphs:
                # Handle image paragraph
                if (p.tag == "img" or getattr(p, "image_path", "")) and p.image_path and os.path.exists(p.image_path):
                    try:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(14)
                        p_img.paragraph_format.space_after = Pt(6)
                        p_img.add_run().add_picture(p.image_path, width=Inches(5.5))
                    except Exception:
                        pass
                    continue

                trans = p.translated_text.strip() if p.translated_text.strip() else p.original_text
                trans = format_math_for_docx(trans)

                if bilingual:
                    # English original
                    en_text = format_math_for_docx(p.original_text)
                    en_p = doc.add_paragraph()
                    en_run = en_p.add_run(en_text)
                    en_run.font.name = "Times New Roman"
                    en_run.font.size = Pt(10)
                    en_run.font.italic = True
                    en_run.font.color.rgb = RGBColor(110, 120, 135)

                    # Vietnamese translation
                    vi_p = doc.add_paragraph()
                    vi_run = vi_p.add_run(trans)
                    vi_run.font.name = "Times New Roman"
                    vi_run.font.size = Pt(11.5)
                    vi_p.paragraph_format.space_after = Pt(8)
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(trans)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    para.paragraph_format.line_spacing = 1.3
                    para.paragraph_format.space_after = Pt(6)

            doc.add_page_break()

        doc.save(output_path)
        return output_path

    @classmethod
    def export_html(cls, project: BookProject, output_path: str, bilingual: bool = False) -> str:
        """Exports as a standalone, printable, elegant HTML reader with MathML and Vietnamese typography."""
        import base64
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

        chapters_html = []
        for chap in project.chapters:
            chap_body = []
            for p in chap.paragraphs:
                # Handle image paragraph (embed base64 so HTML is 100% standalone and printable to PDF)
                if (p.tag == "img" or getattr(p, "image_path", "")) and p.image_path and os.path.exists(p.image_path):
                    try:
                        with open(p.image_path, "rb") as f_img:
                            b64 = base64.b64encode(f_img.read()).decode("utf-8")
                        mime = "image/png" if p.image_path.lower().endswith(".png") else "image/jpeg"
                        chap_body.append(f'''
                        <div class="figure-wrapper" style="text-align: center; margin: 24px auto;">
                            <img src="data:{mime};base64,{b64}" style="max-width: 95%; height: auto; border-radius: 6px; box-shadow: 0 4px 14px rgba(0,0,0,0.12);" />
                        </div>''')
                    except Exception:
                        pass
                    continue

                trans = p.translated_text.strip() if p.translated_text.strip() else p.original_text
                trans_html, _ = format_math_in_html(trans)

                if bilingual:
                    orig_html, _ = format_math_in_html(p.original_text)
                    chap_body.append(f'''
                    <div class="para-pair">
                        <div class="en">{orig_html}</div>
                        <div class="vi">{trans_html}</div>
                    </div>''')
                else:
                    tag = p.tag if p.tag in ('h1', 'h2', 'h3', 'blockquote') else 'p'
                    if trans_html.startswith('<div class="math-block'):
                        parts = trans_html.split('\n', 1)
                        if len(parts) == 2 and parts[1].strip():
                            chap_body.append(parts[0])
                            chap_body.append(f'<{tag}>{parts[1].strip()}</{tag}>')
                        else:
                            chap_body.append(trans_html)
                    else:
                        chap_body.append(f'<{tag}>{trans_html}</{tag}>')

            chapters_html.append(f'''
            <section class="chapter-section" id="{chap.id}">
                <h2>{normalize_text(chap.title)}</h2>
                <div class="chapter-content">
                    {"".join(chap_body)}
                </div>
            </section>
            ''')

        full_html = f'''<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{normalize_text(project.title)} - Bản Dịch</title>
    <link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
    <style>
        :root {{
            --bg-color: #fcfbf9;
            --text-color: #1a202c;
            --accent-color: #2b6cb0;
            --muted-color: #64748b;
            --border-color: #e2e8f0;
        }}
        @media (prefers-color-scheme: dark) {{
            :root {{
                --bg-color: #12141a;
                --text-color: #e2e8f0;
                --accent-color: #818cf8;
                --muted-color: #94a3b8;
                --border-color: #2d3748;
            }}
        }}
        body {{
            background: var(--bg-color);
            color: var(--text-color);
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Times New Roman", "Palatino Linotype", Arial, sans-serif;
            line-height: 1.85;
            margin: 0;
            padding: 40px 20px;
            text-rendering: optimizeLegibility;
            -webkit-font-smoothing: antialiased;
        }}
        .book-container {{
            max-width: 800px;
            margin: 0 auto;
        }}
        header.book-header {{
            text-align: center;
            padding: 60px 0 40px;
            border-bottom: 2px solid var(--border-color);
            margin-bottom: 60px;
            font-family: 'Plus Jakarta Sans', sans-serif;
        }}
        h1.book-title {{
            font-size: 2.5rem;
            margin-bottom: 12px;
            font-weight: 700;
            color: var(--text-color);
        }}
        .book-author {{
            font-size: 1.2rem;
            color: var(--muted-color);
        }}
        .chapter-section {{
            margin-bottom: 80px;
            page-break-after: always;
        }}
        h2 {{
            font-family: 'Plus Jakarta Sans', sans-serif;
            font-size: 1.8rem;
            color: var(--accent-color);
            margin-bottom: 30px;
            text-align: center;
        }}
        p {{
            margin-bottom: 1.2em;
            text-indent: 1.8em;
            text-align: justify;
        }}
        .para-pair {{
            margin-bottom: 24px;
            padding: 12px 16px;
            background: rgba(125, 125, 125, 0.05);
            border-radius: 8px;
            border-left: 3px solid var(--accent-color);
        }}
        .para-pair .en {{
            color: var(--muted-color);
            font-size: 0.9em;
            font-style: italic;
            margin-bottom: 8px;
        }}
        .para-pair .vi {{
            font-size: 1.05em;
        }}
        .math-block {{
            text-align: center;
            margin: 1.4em 0;
            padding: 12px 16px;
            background: rgba(0, 0, 0, 0.03);
            border-radius: 6px;
            overflow-x: auto;
        }}
        .math-equation {{
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        .math-formula {{
            display: inline-block;
            margin: 0 auto;
        }}
        .eq-num {{
            float: right;
            color: var(--muted-color);
            font-size: 0.9em;
        }}
        math {{
            font-family: "Cambria Math", "Latin Modern Math", "STIX Two Math", serif;
            font-size: 1.05em;
        }}
        .math-fallback {{
            font-family: "Cambria Math", "Latin Modern Math", "Times New Roman", serif;
            font-style: italic;
        }}
        @media print {{
            body {{ padding: 0; background: #fff; color: #000; }}
            .chapter-section {{ page-break-after: always; }}
            .para-pair {{ background: transparent; border-left: 1px solid #ccc; }}
        }}
    </style>
</head>
<body>
    <div class="book-container">
        <header class="book-header">
            <h1 class="book-title">{normalize_text(project.title)}</h1>
            <div class="book-author">Tác giả: {normalize_text(project.author)}</div>
            <div style="margin-top: 10px; color: var(--muted-color); font-size: 0.9em;">
                {"Bản dịch Song Ngữ Anh - Việt" if bilingual else "Bản dịch Tiếng Việt (AI Literary Translation)"}
            </div>
        </header>
        {"".join(chapters_html)}
    </div>
</body>
</html>'''

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        return output_path

    @classmethod
    def export_txt(cls, project: BookProject, output_path: str, bilingual: bool = False) -> str:
        """Exports as plain UTF-8 text."""
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        lines = [
            f"=== {project.title} ===",
            f"Tác giả: {project.author}",
            "=" * 40,
            ""
        ]

        for chap in project.chapters:
            lines.append(f"\n\n--- {chap.title} ---\n")
            for p in chap.paragraphs:
                trans = p.translated_text.strip() if p.translated_text.strip() else p.original_text
                if bilingual:
                    lines.append(f"[EN] {p.original_text}")
                    lines.append(f"[VI] {trans}\n")
                else:
                    lines.append(f"{trans}\n")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))
        return output_path
