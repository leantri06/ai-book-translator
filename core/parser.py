"""
Multi-format Book Parser (EPUB, PDF, DOCX, TXT, Markdown).
Extracts structured chapters, paragraphs, headings, and metadata while preserving formatting cues.
"""
from dataclasses import dataclass, field, asdict
from typing import List, Optional, Dict, Tuple
import os
import re
import uuid
from bs4 import BeautifulSoup, NavigableString, Tag
import ebooklib
from ebooklib import epub
import pypdf
import docx


@dataclass
class BookParagraph:
    id: str
    original_text: str
    translated_text: str = ""
    status: str = "pending"  # pending, translating, done, edited, error
    tag: str = "p"           # p, h1, h2, h3, blockquote, li, etc.
    index: int = 0
    css_class: str = ""
    notes: str = ""


@dataclass
class BookChapter:
    id: str
    title: str
    paragraphs: List[BookParagraph] = field(default_factory=list)
    doc_name: str = ""       # e.g., 'chapter1.xhtml' inside EPUB
    order: int = 0
    raw_html: str = ""

    @property
    def total_paragraphs(self) -> int:
        return len(self.paragraphs)

    @property
    def translated_paragraphs(self) -> int:
        return sum(1 for p in self.paragraphs if p.status in ("done", "edited"))

    @property
    def progress_percent(self) -> float:
        if not self.paragraphs:
            return 100.0
        return round((self.translated_paragraphs / len(self.paragraphs)) * 100, 1)

    @property
    def word_count(self) -> int:
        return sum(len(p.original_text.split()) for p in self.paragraphs)


@dataclass
class BookProject:
    id: str
    title: str
    author: str = "Unknown"
    source_format: str = "epub"
    source_file_path: str = ""
    cover_image_path: str = ""
    chapters: List[BookChapter] = field(default_factory=list)
    created_at: str = ""
    updated_at: str = ""

    @property
    def total_chapters(self) -> int:
        return len(self.chapters)

    @property
    def total_paragraphs(self) -> int:
        return sum(c.total_paragraphs for c in self.chapters)

    @property
    def translated_paragraphs(self) -> int:
        return sum(c.translated_paragraphs for c in self.chapters)

    @property
    def total_words(self) -> int:
        return sum(c.word_count for c in self.chapters)

    @property
    def progress_percent(self) -> float:
        total = self.total_paragraphs
        if total == 0:
            return 0.0
        return round((self.translated_paragraphs / total) * 100, 1)


class BookParser:
    """Detects format and parses book files into structured chapters and paragraphs."""

    @staticmethod
    def parse_file(file_path: str, project_id: Optional[str] = None) -> BookProject:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        pid = project_id or str(uuid.uuid4())[:8]

        if ext == ".epub":
            return BookParser.parse_epub(file_path, pid)
        elif ext == ".pdf":
            return BookParser.parse_pdf(file_path, pid)
        elif ext in (".docx", ".doc"):
            return BookParser.parse_docx(file_path, pid)
        elif ext in (".txt", ".md"):
            return BookParser.parse_text(file_path, pid)
        else:
            raise ValueError(f"Định dạng file không được hỗ trợ: {ext}. Hãy tải lên file EPUB, PDF, DOCX, TXT.")

    @staticmethod
    def parse_epub(file_path: str, project_id: str) -> BookProject:
        book = epub.read_epub(file_path)

        # Extract title & author
        title_meta = book.get_metadata('DC', 'title')
        title = title_meta[0][0] if title_meta else os.path.splitext(os.path.basename(file_path))[0]
        creator_meta = book.get_metadata('DC', 'creator')
        author = creator_meta[0][0] if creator_meta else "Tác giả không rõ"

        # 1. Build Table of Contents (TOC) Map: {doc_name: chapter_title}
        toc_map: Dict[str, str] = {}

        def extract_toc(toc_list):
            for item in toc_list:
                if isinstance(item, tuple):
                    if hasattr(item[0], 'href') and hasattr(item[0], 'title'):
                        href = item[0].href.split('#')[0]
                        toc_map[href] = item[0].title.strip()
                    if len(item) > 1 and isinstance(item[1], (list, tuple)):
                        extract_toc(item[1])
                elif hasattr(item, 'href') and hasattr(item, 'title'):
                    href = item.href.split('#')[0]
                    toc_map[href] = item.title.strip()

        if hasattr(book, 'toc') and book.toc:
            extract_toc(book.toc)

        # Look for cover image
        cover_path = ""
        cache_dir = os.path.join(os.path.dirname(file_path), ".covers")
        os.makedirs(cache_dir, exist_ok=True)
        for item in book.get_items_of_type(ebooklib.ITEM_IMAGE):
            name_lower = item.get_name().lower()
            if "cover" in name_lower or "front" in name_lower:
                cover_dest = os.path.join(cache_dir, f"{project_id}_cover.jpg")
                with open(cover_dest, "wb") as cf:
                    cf.write(item.get_content())
                cover_path = cover_dest
                break

        chapters: List[BookChapter] = []
        doc_items = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))

        chap_idx = 0
        p_global_idx = 0

        for doc in doc_items:
            doc_name = doc.get_name()
            content_bytes = doc.get_content()
            try:
                html_str = content_bytes.decode('utf-8')
            except UnicodeDecodeError:
                html_str = content_bytes.decode('latin-1', errors='ignore')

            soup = BeautifulSoup(html_str, 'html.parser')

            # Extract paragraphs and headings, including <div> blocks used in Calibre books
            raw_elements = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'blockquote', 'li', 'div'])
            filtered_elements = []
            for el in raw_elements:
                # If it's a div, skip container divs that contain child paragraphs or child divs
                if el.name == 'div' and el.find(['p', 'div', 'blockquote']):
                    continue
                filtered_elements.append(el)

            chapter_paras: List[BookParagraph] = []
            in_doc_title = ""

            for t in filtered_elements:
                text = t.get_text(strip=True)
                # Clean stray replacement / diamond characters
                text = text.replace('\ufffd', '').strip()

                # Ignore empty elements, trivial single characters or pure ornament dots
                if not text or len(text) < 2 or text in ('* * *', '***', '• • •', '---'):
                    continue

                if t.name in ['h1', 'h2', 'h3'] and not in_doc_title:
                    in_doc_title = text

                p_global_idx += 1
                css_cls = " ".join(t.get('class', [])) if t.has_attr('class') else ""
                tag_name = t.name if t.name in ['h1', 'h2', 'h3', 'blockquote'] else 'p'
                chapter_paras.append(BookParagraph(
                    id=f"c{chap_idx}_p{p_global_idx}",
                    original_text=text,
                    tag=tag_name,
                    index=p_global_idx,
                    css_class=css_cls
                ))

            # Only add as a chapter if there is text
            if chapter_paras:
                # Determine best chapter title:
                # Priority 1: From authoritative EPUB Table of Contents (TOC)
                matched_toc = toc_map.get(doc_name) or toc_map.get(os.path.basename(doc_name))
                if matched_toc:
                    chap_title = matched_toc
                elif in_doc_title:
                    chap_title = in_doc_title
                else:
                    # Look at first paragraph: if it starts with "Chapter", "Prologue", etc.
                    first_text = chapter_paras[0].original_text
                    if re.match(r'^(?:chapter|part|prologue|epilogue|chương|hồi)\b', first_text, re.IGNORECASE) and len(first_text) < 100:
                        chap_title = first_text
                    elif len(first_text) <= 50:
                        chap_title = first_text
                    else:
                        chap_title = f"Phần {chap_idx + 1}: " + (first_text[:50] + "...")

                chapters.append(BookChapter(
                    id=f"chap_{chap_idx}",
                    title=chap_title,
                    paragraphs=chapter_paras,
                    doc_name=doc_name,
                    order=chap_idx,
                    raw_html=html_str
                ))
                chap_idx += 1

        return BookProject(
            id=project_id,
            title=title,
            author=author,
            source_format="epub",
            source_file_path=file_path,
            cover_image_path=cover_path,
            chapters=chapters
        )


    @staticmethod
    def parse_pdf(file_path: str, project_id: str) -> BookProject:
        reader = pypdf.PdfReader(file_path)
        meta = reader.metadata
        raw_name = os.path.splitext(os.path.basename(file_path))[0]
        # Clean hash/project_id prefix like "76f1b7f2_Transformer"
        clean_name = re.sub(r'^[0-9a-f]{8}_', '', raw_name)
        title = meta.title if (meta and meta.title and meta.title.strip()) else clean_name
        author = meta.author if (meta and meta.author and meta.author.strip()) else "Tác giả không rõ"

        MAJOR_HEADING_PATTERN = re.compile(
            r'^(?:'
            r'(\d{1,2}\.?\s+[A-Z][\w\s\-/,\(\)]{2,})'
            r'|'
            r'(Abstract|Conclusion|Conclusions|References|Bibliography|Acknowledgements|Appendix(?:\s+[A-Z0-9]+)?)'
            r'|'
            r'((?:Chapter|Chương|Part|Phần|Section|Hồi|Mục)\s+([0-9ivxlc]+|[a-z]+)[:\s\.\-]*(.*))'
            r')$',
            re.IGNORECASE
        )

        SUB_HEADING_PATTERN = re.compile(
            r'^(?:\d+\.)+\d+\s+([A-Z][\w\s\-/,\(\)]{2,})$'
        )

        chapters: List[BookChapter] = []
        p_global_idx = 0
        chap_idx = 0

        current_paras: List[BookParagraph] = []
        current_title = "Phần mở đầu / Tiêu đề"
        current_lines: List[str] = []

        def flush_para(tag: str = "p"):
            nonlocal current_lines, current_paras, p_global_idx
            if not current_lines:
                return
            text = ""
            for l in current_lines:
                l = l.strip()
                if not l:
                    continue
                if text.endswith("-"):
                    text = text[:-1] + l
                else:
                    text = (text + " " + l).strip() if text else l
            current_lines = []
            if len(text) >= 2:
                p_global_idx += 1
                current_paras.append(BookParagraph(
                    id=f"c{chap_idx}_p{p_global_idx}",
                    original_text=text,
                    tag=tag,
                    index=p_global_idx
                ))

        def flush_chapter(next_title: str):
            nonlocal current_paras, chapters, current_title, chap_idx
            flush_para()
            if current_paras:
                chapters.append(BookChapter(
                    id=f"chap_{chap_idx}",
                    title=current_title,
                    paragraphs=current_paras,
                    doc_name=f"Section_{chap_idx}",
                    order=chap_idx
                ))
                chap_idx += 1
                current_paras = []
            current_title = next_title

        # Check if first page contains paper title (e.g. line in title case before abstract)
        if reader.pages:
            first_page_text = reader.pages[0].extract_text() or ""
            f_lines = [l.strip() for l in first_page_text.splitlines() if l.strip()]
            for l in f_lines[:8]:
                if "attribution" in l.lower() or "permission" in l.lower() or "arxiv" in l.lower():
                    continue
                if 10 < len(l) < 80 and not l.endswith(('.', ':', ';', '@')):
                    if title == clean_name or title == "Tác giả không rõ":
                        title = l
                    break

        for page_num, page in enumerate(reader.pages):
            page_raw = page.extract_text() or ""
            raw_lines = [l.strip() for l in page_raw.splitlines() if l.strip()]
            if not raw_lines:
                continue

            # Filter standalone page number footer at bottom of page
            if raw_lines and re.match(r'^\d+$', raw_lines[-1]):
                raw_lines.pop()

            # Compute typical line length for justified paragraphs on this page
            long_lines = [len(l) for l in raw_lines if len(l) > 30 and not l.endswith(('.', ':', ';'))]
            avg_line_len = (sum(long_lines) / len(long_lines)) if long_lines else 80

            for l_idx, line in enumerate(raw_lines):
                # Check major heading (Starts a new Chapter/Section in TOC)
                m_major = MAJOR_HEADING_PATTERN.match(line)
                if m_major and len(line) < 75 and not line.endswith(('.', ',', ';')):
                    flush_chapter(line)
                    continue

                # Check sub heading (e.g. 3.1, 3.2.1)
                m_sub = SUB_HEADING_PATTERN.match(line)
                if m_sub and len(line) < 75 and not line.endswith(('.', ',', ';')):
                    flush_para()
                    current_lines.append(line)
                    flush_para(tag="h3")
                    continue

                # Check special items (bullets, figures, tables, reference items [1])
                is_bullet = line.startswith(('•', '–', '- ', '* '))
                is_caption = bool(re.match(r'^(?:Figure|Table)\s+\d+[:\.]', line, re.I))
                is_ref_item = bool(re.match(r'^\[\d+\]\s+[A-Z]', line))

                if is_bullet or is_caption or is_ref_item:
                    flush_para()
                    current_lines.append(line)
                    continue

                current_lines.append(line)

                # End of paragraph detection:
                ends_sentence = line.endswith(('.', '!', '?', ':', '."'))
                is_short_line = len(line) < (avg_line_len * 0.78)

                next_starts_new_block = False
                if l_idx + 1 < len(raw_lines):
                    next_l = raw_lines[l_idx + 1]
                    if (MAJOR_HEADING_PATTERN.match(next_l) or 
                        SUB_HEADING_PATTERN.match(next_l) or 
                        next_l.startswith(('•', '–', '- ', '* ')) or 
                        re.match(r'^(?:Figure|Table)\s+\d+[:\.]', next_l, re.I) or
                        re.match(r'^\[\d+\]\s+[A-Z]', next_l)):
                        next_starts_new_block = True

                if ends_sentence and (is_short_line or next_starts_new_block):
                    flush_para()

        flush_chapter("End")

        return BookProject(
            id=project_id,
            title=title,
            author=author,
            source_format="pdf",
            source_file_path=file_path,
            chapters=chapters
        )

    @staticmethod
    def parse_docx(file_path: str, project_id: str) -> BookProject:
        doc = docx.Document(file_path)
        title = os.path.splitext(os.path.basename(file_path))[0]
        author = "Tác giả không rõ"

        chapters: List[BookChapter] = []
        current_paras: List[BookParagraph] = []
        current_title = "Chương 1"
        chap_idx = 0
        p_global_idx = 0

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            style_name = p.style.name.lower() if p.style else ""
            is_heading = "heading 1" in style_name or "title" in style_name or re.match(r'^(?:chapter|chương)\s+\d+', text, re.I)

            if is_heading and current_paras:
                chapters.append(BookChapter(
                    id=f"chap_{chap_idx}",
                    title=current_title,
                    paragraphs=current_paras,
                    order=chap_idx
                ))
                chap_idx += 1
                current_paras = []
                current_title = text

            p_global_idx += 1
            tag = "h1" if is_heading else ("h2" if "heading" in style_name else "p")
            current_paras.append(BookParagraph(
                id=f"c{chap_idx}_p{p_global_idx}",
                original_text=text,
                tag=tag,
                index=p_global_idx
            ))

        if current_paras:
            chapters.append(BookChapter(
                id=f"chap_{chap_idx}",
                title=current_title,
                paragraphs=current_paras,
                order=chap_idx
            ))

        return BookProject(
            id=project_id,
            title=title,
            author=author,
            source_format="docx",
            source_file_path=file_path,
            chapters=chapters
        )

    @staticmethod
    def parse_text(file_path: str, project_id: str) -> BookProject:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            full_text = f.read()

        title = os.path.splitext(os.path.basename(file_path))[0]
        author = "Tác giả không rõ"

        chapter_header_pattern = re.compile(
            r'^(?:#+\s*|CHAPTER\s+|CHƯƠNG\s+|Part\s+)([0-9ivxlc]+|[a-z]+)[:\s\.\-]*(.*)$',
            re.IGNORECASE | re.MULTILINE
        )

        lines = full_text.splitlines()
        chapters: List[BookChapter] = []
        current_paras: List[BookParagraph] = []
        current_title = "Phần 1"
        chap_idx = 0
        p_global_idx = 0

        # Buffer for continuous paragraph lines
        buf = []

        def flush_buffer(target_tag="p"):
            nonlocal p_global_idx, buf
            if buf:
                p_text = " ".join(buf).strip()
                if p_text:
                    p_global_idx += 1
                    current_paras.append(BookParagraph(
                        id=f"c{chap_idx}_p{p_global_idx}",
                        original_text=p_text,
                        tag=target_tag,
                        index=p_global_idx
                    ))
                buf = []

        for line in lines:
            line_str = line.strip()
            if not line_str:
                flush_buffer()
                continue

            match = chapter_header_pattern.match(line_str)
            if match and len(line_str) < 100:
                flush_buffer()
                if current_paras:
                    chapters.append(BookChapter(
                        id=f"chap_{chap_idx}",
                        title=current_title,
                        paragraphs=current_paras,
                        order=chap_idx
                    ))
                    chap_idx += 1
                    current_paras = []
                current_title = line_str.lstrip('#').strip()
                p_global_idx += 1
                current_paras.append(BookParagraph(
                    id=f"c{chap_idx}_p{p_global_idx}",
                    original_text=current_title,
                    tag="h1",
                    index=p_global_idx
                ))
            else:
                buf.append(line_str)

        flush_buffer()
        if current_paras:
            chapters.append(BookChapter(
                id=f"chap_{chap_idx}",
                title=current_title,
                paragraphs=current_paras,
                order=chap_idx
            ))

        return BookProject(
            id=project_id,
            title=title,
            author=author,
            source_format="txt",
            source_file_path=file_path,
            chapters=chapters
        )
